from openpyxl import load_workbook
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os

# =========================
# RUTAS DE ARCHIVOS
# =========================
ruta_excel = os.path.expanduser("C:\Users\jimgl\Downloads\RUTA_PROCESADA.xlsx")
ruta_word = os.path.expanduser("C:\Users\jimgl\Downloads\RUTA_PROCESADA.docx")

# =========================
# CARGAR EXCEL
# =========================
wb = load_workbook(ruta_excel, data_only=True)
ws = wb["RUTA"]

# Leer todas las filas con datos
rows = list(ws.iter_rows(values_only=True))
max_cols = max(len(row) for row in rows)

# =========================
# CREAR DOCUMENTO WORD
# =========================
doc = Document()

# -------------------------
# CONFIGURACIÓN DE PÁGINA
# -------------------------
section = doc.sections[0]
section.orientation = WD_ORIENT.LANDSCAPE
section.page_width, section.page_height = section.page_height, section.page_width

# Márgenes ultra reducidos
section.top_margin = Cm(0.7)
section.bottom_margin = Cm(0.7)
section.left_margin = Cm(0.6)
section.right_margin = Cm(0.6)

# =========================
# FUNCIÓN: QUITAR PADDING DE CELDAS
# =========================
def remove_table_cell_margins(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblCellMar = OxmlElement('w:tblCellMar')

    for margin in ['top', 'left', 'bottom', 'right']:
        node = OxmlElement(f'w:{margin}')
        node.set(qn('w:w'), '0')
        node.set(qn('w:type'), 'dxa')
        tblCellMar.append(node)

    tblPr.append(tblCellMar)

# =========================
# CREAR TABLA
# =========================
table = doc.add_table(rows=len(rows), cols=max_cols)
table.style = "Table Grid"
table.autofit = True

remove_table_cell_margins(table)

# =========================
# LLENAR TABLA (COMPACTO)
# =========================
for i, row in enumerate(rows):
    for j, cell in enumerate(row):
        if cell is not None:
            paragraph = table.rows[i].cells[j].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1

            run = paragraph.add_run(str(cell))
            run.font.size = Pt(8.5)

# =========================
# EVITAR DIVISIÓN DE FILAS ENTRE PÁGINAS
# =========================
for row in table.rows:
    for cell in row.cells:
        tcPr = cell._tc.get_or_add_tcPr()
        noWrap = OxmlElement('w:noWrap')
        tcPr.append(noWrap)

# =========================
# NUMERACIÓN DE PÁGINAS
# =========================
footer = section.footer
p = footer.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

run = p.add_run()
fldChar1 = OxmlElement('w:fldChar')
fldChar1.set(qn('w:fldCharType'), 'begin')

instrText = OxmlElement('w:instrText')
instrText.text = "PAGE"

fldChar2 = OxmlElement('w:fldChar')
fldChar2.set(qn('w:fldCharType'), 'end')

run._r.append(fldChar1)
run._r.append(instrText)
run._r.append(fldChar2)

# =========================
# GUARDAR DOCUMENTO
# =========================
doc.save(ruta_word)

print("✅ Documento Word generado y optimizado correctamente")
